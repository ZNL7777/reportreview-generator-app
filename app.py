import streamlit as st
import pandas as pd
import docx
import io
import zipfile
import re
from datetime import datetime

def clean_empty_paragraphs_and_runs(doc):
    """
    清理 Word 文档正文及表格单元格中的隐形空行、多余空段落和多余换行符
    """
    for p in list(doc.paragraphs):
        if not p.text.strip() and len(doc.paragraphs) > 1:
            if not p._element.xpath('.//w:drawing') and not p._element.xpath('.//w:object'):
                p._element.getparent().remove(p._element)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs = list(cell.paragraphs)
                for p in paragraphs:
                    if len(cell.paragraphs) > 1 and not p.text.strip():
                        if not p._element.xpath('.//w:drawing') and not p._element.xpath('.//w:object'):
                            p._element.getparent().remove(p._element)
                
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.text:
                            run.text = run.text.strip('\r\n')

def process_single_row(doc, row):
    """
    对单行数据进行 Word 模板匹配填充及精确勾选逻辑判断
    """
    # 提取 Excel 行数据
    task_no = str(row.get('任务号', '') or '').strip()
    company_name = str(row.get('公司名称', '') or '').strip()
    
    # 审核组长：取第一个名字
    auditor_raw = str(row.get('审核组长', '') or '').strip()
    auditor_leader = re.split(r'[,，\s/+]+', auditor_raw)[0] if auditor_raw else ''
    
    audit_address = str(row.get('审核地址', '') or '').strip()
    cert_scope = str(row.get('认证范围', '') or '').strip()
    
    # 日期格式化处理
    raw_date = row.get('日期', '')
    if pd.notna(raw_date):
        if isinstance(raw_date, (int, float)):
            date_str = (datetime(1899, 12, 30) + pd.Timedelta(days=raw_date)).strftime('%Y-%m-%d')
        else:
            try:
                date_str = pd.to_datetime(raw_date).strftime('%Y-%m-%d')
            except:
                date_str = str(raw_date).strip()
    else:
        date_str = ''
    
    audit_type = str(row.get('审核类型Audit Type', '') or row.get('审核类型', '') or '').strip()
    cert_decision = str(row.get('认证决定结论', '') or '').strip()
    
    # --- 逻辑条件计算 ---
    task_upper = task_no.upper()
    has_ts = 'TS' in task_upper
    has_er = 'ER' in task_upper
    
    chk_ts = has_ts and not has_er
    chk_er = has_er and not has_ts
    chk_both = has_ts and has_er
    
    chk_stage2 = ('二阶段' in audit_type) or ('初审' in audit_type)
    chk_surveillance = ('监' in audit_type) or ('监一' in audit_type) or ('监二' in audit_type)
    chk_recert_transfer = ('再认证' in audit_type) or ('转移' in audit_type)
    chk_special = '特殊' in audit_type
    
    # 认证决定结论判断
    chk_dec_1 = chk_stage2 or ('再认证' in audit_type)
    chk_dec_3 = '转移' in audit_type
    chk_dec_4 = chk_special and ('换发' in cert_decision)
    chk_dec_5 = chk_surveillance and ('不换证' in cert_decision or '保持' in cert_decision)
    chk_dec_6 = chk_surveillance and ('换发' in cert_decision)
    
    # --- 替换辅助函数 ---
    def replace_text_in_paragraph(p):
        text = p.text
        if not text:
            return
            
        original_text = text

        # 基础占位符替换
        placeholders = {
            r'\{\{\s*任务号\s*\}\}': task_no,
            r'\{\{\s*公司名称\s*\}\}': company_name,
            r'\{\{\s*审核组长\s*\}\}': auditor_leader,
            r'\{\{\s*审核地址\s*\}\}': audit_address,
            r'\{\{\s*认证范围\s*\}\}': cert_scope,
            r'\{\{\s*日期\s*\}\}': date_str,
            
            r'\{\{\s*iatf_check\s*\}\}': '☑' if (chk_ts or chk_both) else '□',
            r'\{\{\s*iso_check\s*\}\}': '☑' if (chk_er or chk_both) else '□',
            
            r'\{\{\s*chu_shen\s*\}\}': '☑' if chk_stage2 else '□',
            r'\{\{\s*jian_shen\s*\}\}': '☑' if chk_surveillance else '□',
            r'\{\{\s*zai_ren_zheng\s*\}\}': '☑' if chk_recert_transfer else '□',
            r'\{\{\s*te_shu\s*\}\}': '☑' if chk_special else '□',
            
            # 认证决定结论：选中时显示带 X 的大方框形式，未选中时显示标准大方框
            r'\{\{\s*dec_1\s*\}\}': '☒' if chk_dec_1 else '□',
            r'\{\{\s*dec_3\s*\}\}': '☒' if chk_dec_3 else '□',
            r'\{\{\s*dec_4\s*\}\}': '☒' if chk_dec_4 else '□',
            r'\{\{\s*dec_5\s*\}\}': '☒' if chk_dec_5 else '□',
            r'\{\{\s*dec_6\s*\}\}': '☒' if chk_dec_6 else '□',
        }

        for pattern, value in placeholders.items():
            text = re.sub(pattern, str(value), text)
            
        # 针对直接写在模板中的原生大方框进行统一匹配替换（选中时打X即 ☒，未选中时为空方框 □）
        if '通过，可发证（适用于：初审' in text:
            box = '☒' if chk_dec_1 else '□'
            text = re.sub(r'^[☑□☒\s]*', box + ' ', text)
        elif '通过，暂停恢复审核' in text:
            text = re.sub(r'^[☑□☒\s]*', '□ ', text)
        elif '通过，可换发证书（转机构）' in text:
            box = '☒' if chk_dec_3 else '□'
            text = re.sub(r'^[☑□☒\s]*', box + ' ', text)
        elif '通过，同意扩大认证范围' in text:
            box = '☒' if chk_dec_4 else '□'
            text = re.sub(r'^[☑□☒\s]*', box + ' ', text)
        elif '通过，不换证' in text:
            box = '☒' if chk_dec_5 else '□'
            text = re.sub(r'^[☑□☒\s]*', box + ' ', text)
        elif '通过，可换发新的认证证书' in text:
            box = '☒' if chk_dec_6 else '□'
            text = re.sub(r'^[☑□☒\s]*', box + ' ', text)
        elif '不予通过' in text:
            text = re.sub(r'^[☑□☒\s]*', '□ ', text)

        if text != original_text:
            p.text = text

    for p in doc.paragraphs:
        replace_text_in_paragraph(p)
        
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_text_in_paragraph(p)

    clean_empty_paragraphs_and_runs(doc)

# --- Streamlit 界面构建 ---
st.set_page_config(page_title="Word 批量生成与空行清理工具", layout="centered")
st.title("📄 Word 批量生成与空行清理工具")
st.write("上传 Excel 表格和 Word 模板，系统将自动填充数据、统一规范复选框（选中打X、未选为空方框）、清理多余空行后打包下载。")

uploaded_excel = st.file_uploader("上传 Excel 数据文件 (.xlsx / .xls)", type=["xlsx", "xls"])
uploaded_template = st.file_uploader("上传 Word 模板文件 (.docx)", type=["docx"])

if uploaded_excel and uploaded_template:
    try:
        df_raw = pd.read_excel(uploaded_excel)
        df = df_raw.dropna(how='all').fillna('')
        if '任务号' in df.columns and '公司名称' in df.columns:
            df = df[~((df['任务号'] == '') & (df['公司名称'] == ''))]
            
        st.success(f"成功读取并过滤 Excel 文件，共计保留 {len(df)} 行有效数据（原数据 {len(df_raw)} 行）。")
        st.dataframe(df.head())
        
        if st.button("🚀 开始清理并生成文件"):
            zip_buffer = io.BytesIO()
            template_bytes = uploaded_template.getvalue()
            
            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                for idx, row in df.iterrows():
                    doc = docx.Document(io.BytesIO(template_bytes))
                    process_single_row(doc, row)
                    
                    out_io = io.BytesIO()
                    doc.save(out_io)
                    
                    task_id = str(row.get('任务号', f'Row_{idx+1}')).strip()
                    company_name = str(row.get('公司名称', '')).strip()
                    filename = f"{task_id}_{company_name}.docx" if company_name else f"{task_id}.docx"
                    filename = re.sub(r'[\\/:*?"<>|]', '_', filename)
                    
                    zip_file.writestr(filename, out_io.getvalue())
            
            st.success("🎉 所有 Word 文件已完成生成与隐形空行清理！")
            st.download_button(
                label="📦 点击下载打包 ZIP 文件",
                data=zip_buffer.getvalue(),
                file_name="生成的清理后Word文档集.zip",
                mime="application/zip"
            )
    except Exception as e:
        st.error(f"处理过程中出错: {e}")
